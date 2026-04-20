"""Generate tailored resumes as PDF or DOCX."""
import io
from typing import Literal

def generate_pdf(resume_text: str, filename: str = 'resume') -> bytes:
    """Generate a clean PDF from plain-text resume."""
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.enums import TA_LEFT, TA_CENTER

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=LETTER,
            leftMargin=0.75 * inch,
            rightMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )

        styles = getSampleStyleSheet()
        # Custom styles
        name_style = ParagraphStyle('Name', parent=styles['Title'],
                                     fontSize=18, textColor=colors.HexColor('#1e293b'),
                                     spaceAfter=4, alignment=TA_CENTER)
        heading_style = ParagraphStyle('Heading', parent=styles['Heading2'],
                                        fontSize=11, textColor=colors.HexColor('#4338ca'),
                                        spaceBefore=10, spaceAfter=4, borderPadding=(0,0,2,0))
        body_style = ParagraphStyle('Body', parent=styles['Normal'],
                                     fontSize=9.5, leading=14,
                                     textColor=colors.HexColor('#334155'))
        bullet_style = ParagraphStyle('Bullet', parent=body_style,
                                       leftIndent=12, bulletIndent=0, spaceBefore=1)

        story = []
        lines = resume_text.strip().split('\n')

        for i, line in enumerate(lines):
            stripped = line.strip()
            if not stripped:
                story.append(Spacer(1, 4))
                continue

            # Heuristics for formatting
            if i == 0 and len(stripped) < 60:
                story.append(Paragraph(stripped, name_style))
                story.append(HRFlowable(width='100%', thickness=1, color=colors.HexColor('#4338ca')))
            elif stripped.isupper() or (stripped.endswith(':') and len(stripped) < 30):
                story.append(Spacer(1, 6))
                story.append(Paragraph(stripped.rstrip(':'), heading_style))
                story.append(HRFlowable(width='100%', thickness=0.5, color=colors.HexColor('#e2e8f0')))
            elif stripped.startswith(('•', '-', '*', '·')):
                story.append(Paragraph('• ' + stripped.lstrip('•-*· '), bullet_style))
            else:
                story.append(Paragraph(stripped, body_style))

        doc.build(story)
        return buf.getvalue()

    except ImportError:
        # Fallback: minimal PDF with fpdf2
        return _generate_pdf_fpdf(resume_text)


def _generate_pdf_fpdf(resume_text: str) -> bytes:
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font('Helvetica', size=10)
    for line in resume_text.split('\n'):
        pdf.multi_cell(0, 5, line)
    return pdf.output()


def generate_docx(resume_text: str) -> bytes:
    """Generate a formatted DOCX from plain-text resume."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    lines = resume_text.strip().split('\n')

    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph('')
            continue

        if i == 0 and len(stripped) < 60:
            # Name / title line
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)

        elif stripped.isupper() or (stripped.endswith(':') and len(stripped) < 30):
            # Section heading
            p = doc.add_paragraph()
            run = p.add_run(stripped.rstrip(':'))
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x43, 0x38, 0xca)
            p.paragraph_format.space_before = Pt(8)
            # Bottom border via paragraph border (Word quirk)

        elif stripped.startswith(('•', '-', '*', '·')):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(stripped.lstrip('•-*· '))
            p.paragraph_format.left_indent = Inches(0.2)

        else:
            p = doc.add_paragraph()
            p.add_run(stripped)
            p.paragraph_format.space_after = Pt(2)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
