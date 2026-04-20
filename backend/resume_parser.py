"""Parse PDF / DOCX / TXT resumes into plain text."""
import io
from pathlib import Path

def parse_resume(file_bytes: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == '.pdf':
        return _parse_pdf(file_bytes)
    elif ext in ('.docx', '.doc'):
        return _parse_docx(file_bytes)
    else:
        return file_bytes.decode('utf-8', errors='replace')

def _parse_pdf(data: bytes) -> str:
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(data))
        lines = []
        for page in reader.pages:
            txt = page.extract_text()
            if txt:
                lines.append(txt)
        return '\n'.join(lines)
    except Exception as e:
        raise ValueError(f'PDF parsing failed: {e}')

def _parse_docx(data: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return '\n'.join(parts)
    except Exception as e:
        raise ValueError(f'DOCX parsing failed: {e}')
